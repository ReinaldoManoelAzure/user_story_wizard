import streamlit as st
from dotenv import load_dotenv
import os
import json
import xml.etree.ElementTree as ET
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import ChatPromptTemplate
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from io import BytesIO
from docx import Document
from markdown_pdf import MarkdownPdf, Section

# carregar variáveis do .env
load_dotenv()

st.set_page_config(page_title="User Story Wizard", layout="wide")

# -------------------
# Função para processar arquivos (mini-RAG)
# -------------------
def process_files(uploaded_files):
    conteudo = []
    for file in uploaded_files:
        if file.name.endswith(".pdf"):
            loader = PyPDFLoader(file)
            docs = loader.load()
            conteudo.extend([d.page_content for d in docs])

        elif file.name.endswith(".docx"):
            loader = Docx2txtLoader(file)
            docs = loader.load()
            conteudo.extend([d.page_content for d in docs])

        elif file.name.endswith(".txt"):
            text = file.read().decode("utf-8")
            conteudo.append(text)

        elif file.name.endswith(".json"):
            text = json.dumps(json.load(file), indent=2, ensure_ascii=False)
            conteudo.append(text)

        elif file.name.endswith(".xml"):
            tree = ET.parse(file)
            root = tree.getroot()
            text = ET.tostring(root, encoding="utf-8").decode("utf-8")
            conteudo.append(text)

    return "\n\n".join(conteudo)


# -------------------
# Função para criar a história com fallback de modelos
# -------------------
def build_with_fallback(contexto_texto, macro, inputs, arquivos_contexto):
    prompt = ChatPromptTemplate.from_template("""
    Construa uma **História Funcional** no seguinte formato:

    História Funcional

    Título: <Título claro e objetivo>

    Como
    <persona ou papel do usuário>

    Quero
    <objetivo ou ação desejada>

    Para
    <benefício ou valor agregado>

    Critérios de Aceite
    - Liste critérios claros e verificáveis

    Regras de Negócio
    - Liste as regras de negócio aplicáveis

    Informações Técnicas
    **Itens Prioritários (informados pelo usuário em 'Informações Técnicas ou Observações Extras'):**
    - Liste diretamente os itens fornecidos pelo usuário. 
    - Se o usuário não informou nada, escreva "Nenhum informado."

    **Propostas Adicionais (sugeridas pelo modelo a partir de contexto e arquivos):**
    - Sugira pontos técnicos adicionais que possam ser relevantes. 
    - Essas propostas devem ser opcionais, não obrigatórias.

    ---  
    **Contexto Macro:**  
    {macro}

    **Contexto Detalhado:**  
    {contexto_texto}

    **Inputs adicionais:**  
    {inputs}

    **Informações extraídas de arquivos:**  
    {arquivos_contexto}
    """)

    # 🔄 Fallback: Gemini -> OpenAI
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            temperature=0.1,
            api_key=os.getenv("GOOGLE_API_KEY")
        )
        chain = prompt | llm
        response = chain.invoke({
            "macro": macro,
            "contexto_texto": contexto_texto,
            "inputs": inputs,
            "arquivos_contexto": arquivos_contexto
        })
        return f"{response.content}"

    except Exception as e1:
        try:
            llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.1)
            chain = prompt | llm
            response = chain.invoke({
                "macro": macro,
                "contexto_texto": contexto_texto,
                "inputs": inputs,
                "arquivos_contexto": arquivos_contexto
            })
            return f"(⚠️ Gemini falhou, fallback para OpenAI GPT-4o-mini)\n\n{response.content}"
        except Exception as e2:
            return f"❌ Nenhum modelo pôde ser utilizado.\nErro Gemini: {str(e1)}\nErro OpenAI: {str(e2)}"


# -------------------
# Export helpers
# -------------------
def export_docx(story_text: str) -> BytesIO:
    doc = Document()
    for line in story_text.split("\n"):
        if line.startswith("Título:"):
            doc.add_heading(line.replace("Título:", "").strip(), level=1)
        elif line.startswith("Como"):
            doc.add_heading("Como", level=2)
            doc.add_paragraph(line.replace("Como", "").strip())
        elif line.startswith("Quero"):
            doc.add_heading("Quero", level=2)
            doc.add_paragraph(line.replace("Quero", "").strip())
        elif line.startswith("Para"):
            doc.add_heading("Para", level=2)
            doc.add_paragraph(line.replace("Para", "").strip())
        elif line.strip().endswith(":"):
            doc.add_heading(line.strip(), level=2)
        elif line.strip().startswith("-"):
            doc.add_paragraph(line.strip(), style="List Bullet")
        else:
            doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_pdf(story_text: str) -> BytesIO:
    buffer = BytesIO()
    pdf = MarkdownPdf()
    pdf.add_section(Section(story_text, toc=False))
    pdf.save(buffer)
    buffer.seek(0)
    return buffer


# -------------------
# UI do Streamlit
# -------------------
def main():
    st.title("🧙 User Story Wizard")
    st.markdown("Monte histórias de usuário de forma padronizada.")

    # 📖 Explicação educativa
    with st.expander("ℹ️ O que é uma História de Usuário?"):
        st.markdown("""
        Uma **História de Usuário** é uma forma simples e clara de descrever uma necessidade de negócio do ponto de vista do usuário final.  
        
        Ela geralmente segue o formato:

        **Como [persona]**  
        **Quero [ação ou objetivo]**  
        **Para [benefício ou valor]**
        """)

    st.subheader("🌍 Qual é o cenário geral do projeto?")
    macro = st.text_area(
        "Objetivo geral do sistema",
        height=120,
        placeholder="Sistema de gestão de contratos para empresas de logística, com integração ao ERP."
    )

    st.subheader("📖 Detalhe o que precisa ser feito")
    contexto_texto = st.text_area(
        "Requisitos e descrições do cliente",
        height=200,
        placeholder="""Usuários precisam visualizar contratos ativos e vencidos, receber notificações automáticas de renovação e gerar relatórios detalhados por cliente e filial. 
Gestores administrativos terão acesso a relatórios financeiros consolidados."""
    )

    st.subheader("📎 Anexe documentos de apoio")
    arquivos = st.file_uploader(
        "Envie arquivos (docx, pdf, txt, xml, json)",
        type=["docx", "pdf", "txt", "xml", "json"],
        accept_multiple_files=True
    )
    
    arquivos_contexto = ""
    if arquivos:
        arquivos_contexto = process_files(arquivos)

    st.subheader("🔧 Informações Técnicas ou Observações Extras")
    inputs = st.text_area(
        "Digite restrições técnicas, regras específicas ou observações adicionais",
        height=100,
        placeholder="Compatível com Chrome e Edge.\nAPI deve responder em até 2 segundos.",
        help="Se você não preencher nada, o modelo ainda sugerirá **Propostas Adicionais** automaticamente."
    )
    st.caption("💡 Exemplo: 'Compatível com Chrome e Edge. API deve responder em até 2 segundos.'")

    # -------------------
    # Geração da história
    # -------------------
    if st.button("✨ Criar História Funcional"):
        if not macro and not contexto_texto and not arquivos_contexto:
            st.warning("Por favor, preencha algum campo de contexto ou envie arquivos.")
        else:
            with st.spinner("Gerando história..."):
                st.session_state["story"] = build_with_fallback(
                    contexto_texto, macro, inputs, arquivos_contexto
                )
            st.success("✅ História gerada com sucesso!")

    # -------------------
    # Exibir resultado + botões de download
    # -------------------
    if "story" in st.session_state and st.session_state["story"]:
        story = st.session_state["story"]
        st.markdown("### 📖 Resultado")
        st.markdown(story)

        st.download_button(
            label="📥 Baixar DOCX",
            data=export_docx(story),
            file_name="historia_funcional.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.download_button(
            label="📥 Baixar PDF",
            data=export_pdf(story),
            file_name="historia_funcional.pdf",
            mime="application/pdf"
        )


if __name__ == "__main__":
    main()
